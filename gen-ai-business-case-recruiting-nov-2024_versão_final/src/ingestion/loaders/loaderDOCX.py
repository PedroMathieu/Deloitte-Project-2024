from src.ingestion.loaders.loaderBase import LoaderBase
from docx import Document
from docx.shared import Inches
import pytesseract
from PIL import Image
import io

class LoaderDOCX(LoaderBase):

    def __init__(self, filepath: str):
        self.filepath = filepath

    def extract_metadata(self):
        # Load the .docx file
        doc = Document(self.filepath)
        
        # Extract metadata from core properties
        metadata = {
            'author': doc.core_properties.author,
            'title': doc.core_properties.title,
            'subject': doc.core_properties.subject,
            'keywords': doc.core_properties.keywords,
            'created': doc.core_properties.created,
        }

        self.metadata = metadata
        
        # Check if all metadata values are available (non-empty)
        return self.metadata if self.all_keys_have_values(metadata=self.metadata) else False
    
    def extract_text(self):
        # Load the .docx file
        doc = Document(self.filepath)
        
        # Extract all text from paragraphs and return as a single string
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"  # Adding newlines between paragraphs
        
        return text

    def extract_image_text(self, page_num: int = None):
        """Extract text from images embedded in the .docx file using OCR."""
        doc = Document(self.filepath)
        text_from_images = ""

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref.lower():
                # Open image from .docx and apply OCR
                img_stream = io.BytesIO(rel.target_part._blob)
                image = Image.open(img_stream)
                text_from_images += pytesseract.image_to_string(image)

        return text_from_images

    def extract_chunks(self):
        """Chunk the extracted text into paragraphs."""
        text = self.extract_text()
        
        if not text:
            return []
        
        # Split text into paragraphs by single newline
        paragraphs = text.split("\n\n")  # Split by single newline for paragraphs
        
        # Remove any empty or whitespace-only paragraphs
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        #text_for_splitter = "\n\n".join(paragraphs)
        
        return paragraphs


    def extract_image_chunks(self):
        """Extract text from images in the .docx file and chunk them into paragraphs."""
        text_from_images = self.extract_image_text()
        
        if not text_from_images:
            return []
        
        # Split text from images into paragraphs
        image_chunks = text_from_images.split("\n\n")
        return [chunk.strip() for chunk in image_chunks if chunk.strip()]

    def all_keys_have_values(self, metadata, value_check=lambda x: x is not None and x != ''):
        # Helper function to ensure all metadata values are non-empty
        return all(value_check(value) for value in metadata.values())
